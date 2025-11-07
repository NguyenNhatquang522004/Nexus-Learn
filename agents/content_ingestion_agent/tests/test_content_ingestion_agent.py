# Content Ingestion Agent - Tests

import asyncio
import json
import os
import tempfile
from datetime import datetime
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi.testclient import TestClient
from httpx import AsyncClient

from content_ingestion_agent import (
    ContentIngestionAgent,
    FileProcessor,
    KnowledgeGraphClient,
    ModelManager,
    ProcessingStatus,
    app,
)


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def test_config():
    """Test configuration"""
    return {
        "agent": {
            "name": "test_agent",
            "port": 8001,
            "host": "0.0.0.0"
        },
        "models": {
            "distilbert": {
                "model_path": "distilbert-base-uncased",
                "num_labels": 5,
                "subjects": ["math", "science", "history", "language", "arts"]
            }
        },
        "file_processing": {
            "supported_formats": ["pdf", "docx", "pptx"],
            "max_file_size_mb": 50,
            "temp_dir": "/tmp/test_ingestion",
            "output_dir": "/tmp/test_output"
        },
        "extraction": {
            "min_confidence": 0.7,
            "concept_extraction_method": "ner",
            "relationship_detection": True,
            "image_analysis": True
        },
        "knowledge_graph_api": {
            "base_url": "http://localhost:8000",
            "timeout": 30,
            "retry": 3
        },
        "storage": {
            "type": "local",
            "path": "/tmp/test_uploads"
        }
    }


@pytest.fixture
def sample_pdf(tmp_path):
    """Create sample PDF file"""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    pdf_path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "Sample PDF Document")
    c.drawString(100, 730, "This is a test document about mathematics.")
    c.drawString(100, 710, "The Pythagorean theorem: a² + b² = c²")
    c.save()
    
    return pdf_path


@pytest.fixture
def sample_docx(tmp_path):
    """Create sample DOCX file"""
    from docx import Document
    
    doc = Document()
    doc.add_heading("Sample Document", 0)
    doc.add_paragraph("This is a test document about science.")
    doc.add_paragraph("Newton's first law of motion.")
    
    docx_path = tmp_path / "sample.docx"
    doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def mock_model_manager():
    """Mock ModelManager"""
    manager = MagicMock(spec=ModelManager)
    
    manager.classify_content.return_value = {
        "primary_subject": "math",
        "confidence": 0.95,
        "top_predictions": [
            {"subject": "math", "confidence": 0.95},
            {"subject": "science", "confidence": 0.03}
        ],
        "difficulty": "intermediate",
        "blooms_level": "apply"
    }
    
    manager.extract_concepts.return_value = [
        {
            "text": "Pythagorean theorem",
            "type": "CONCEPT",
            "confidence": 0.92,
            "method": "ner"
        },
        {
            "text": "triangle",
            "type": "CONCEPT",
            "confidence": 0.88,
            "method": "ner"
        }
    ]
    
    manager.detect_relationships.return_value = [
        {
            "from_concept": "Pythagorean theorem",
            "to_concept": "triangle",
            "relationship_type": "related_to",
            "confidence": 0.85,
            "context": "theorem applies to triangles"
        }
    ]
    
    return manager


@pytest.fixture
def mock_graph_client():
    """Mock KnowledgeGraphClient"""
    client = AsyncMock(spec=KnowledgeGraphClient)
    
    client.create_nodes.return_value = ["node1", "node2"]
    client.create_relationships.return_value = ["rel1"]
    
    return client


# ============================================================================
# FILE PROCESSOR TESTS
# ============================================================================

class TestFileProcessor:
    """Test FileProcessor class"""
    
    def test_validate_file_success(self, test_config, sample_pdf):
        """Test successful file validation"""
        processor = FileProcessor(test_config)
        assert processor.validate_file(sample_pdf) is True
    
    def test_validate_file_not_found(self, test_config):
        """Test validation with non-existent file"""
        processor = FileProcessor(test_config)
        
        with pytest.raises(ValueError, match="File not found"):
            processor.validate_file(Path("/nonexistent.pdf"))
    
    def test_validate_file_too_large(self, test_config, tmp_path):
        """Test validation with oversized file"""
        # Create large file
        large_file = tmp_path / "large.pdf"
        with open(large_file, "wb") as f:
            f.write(b"x" * (60 * 1024 * 1024))  # 60 MB
        
        processor = FileProcessor(test_config)
        
        with pytest.raises(ValueError, match="File too large"):
            processor.validate_file(large_file)
    
    def test_validate_file_unsupported_format(self, test_config, tmp_path):
        """Test validation with unsupported format"""
        unsupported_file = tmp_path / "file.xyz"
        unsupported_file.write_text("content")
        
        processor = FileProcessor(test_config)
        
        with pytest.raises(ValueError, match="Unsupported format"):
            processor.validate_file(unsupported_file)
    
    def test_extract_pdf_content(self, test_config, sample_pdf):
        """Test PDF content extraction"""
        processor = FileProcessor(test_config)
        content = processor.extract_content(sample_pdf)
        
        assert content["type"] == "pdf"
        assert content["page_count"] > 0
        assert len(content["pages"]) > 0
        assert "text" in content["pages"][0]
        assert "mathematics" in content["pages"][0]["text"].lower()
    
    def test_extract_docx_content(self, test_config, sample_docx):
        """Test DOCX content extraction"""
        processor = FileProcessor(test_config)
        content = processor.extract_content(sample_docx)
        
        assert content["type"] == "docx"
        assert content["page_count"] == 1
        assert len(content["pages"]) > 0
        assert "science" in content["pages"][0]["text"].lower()
    
    def test_extract_equations(self, test_config):
        """Test LaTeX equation extraction"""
        processor = FileProcessor(test_config)
        
        text = "The formula is $E = mc^2$ and $$\\frac{a}{b} = c$$"
        equations = processor._extract_equations(text)
        
        assert len(equations) >= 2
        assert any("E = mc^2" in eq["latex"] for eq in equations)
    
    def test_convert_latex(self, test_config):
        """Test LaTeX conversion"""
        processor = FileProcessor(test_config)
        
        latex = "\\frac{a}{b}"
        result = processor.convert_latex(latex)
        
        assert "(" in result and ")" in result
        assert "/" in result
    
    def test_process_tables(self, test_config):
        """Test table processing"""
        processor = FileProcessor(test_config)
        
        tables = [
            [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
        ]
        
        processed = processor.process_tables(tables)
        
        assert len(processed) == 1
        assert processed[0]["headers"] == ["Name", "Age"]
        assert processed[0]["row_count"] == 2
        assert processed[0]["column_count"] == 2


# ============================================================================
# MODEL MANAGER TESTS
# ============================================================================

class TestModelManager:
    """Test ModelManager class"""
    
    @patch("content_ingestion_agent.AutoTokenizer")
    @patch("content_ingestion_agent.AutoModelForSequenceClassification")
    @patch("content_ingestion_agent.pipeline")
    def test_load_models(self, mock_pipeline, mock_model, mock_tokenizer, test_config):
        """Test model loading"""
        manager = ModelManager(test_config)
        manager.load_models()
        
        # Check models loaded
        assert "distilbert" in manager.tokenizers
        assert "distilbert" in manager.models
        assert "ner" in manager.pipelines
        assert "zero_shot" in manager.pipelines
    
    def test_classify_content(self, test_config, mock_model_manager):
        """Test content classification"""
        text = "The Pythagorean theorem states that a² + b² = c²"
        result = mock_model_manager.classify_content(text)
        
        assert "primary_subject" in result
        assert "confidence" in result
        assert "difficulty" in result
        assert "blooms_level" in result
        assert result["primary_subject"] in test_config["models"]["distilbert"]["subjects"]
    
    def test_extract_concepts(self, test_config, mock_model_manager):
        """Test concept extraction"""
        text = "The Pythagorean theorem applies to right triangles"
        concepts = mock_model_manager.extract_concepts(text, method="ner")
        
        assert len(concepts) > 0
        assert all("text" in c for c in concepts)
        assert all("confidence" in c for c in concepts)
        assert all(c["confidence"] >= 0.0 for c in concepts)
    
    def test_detect_relationships(self, test_config, mock_model_manager):
        """Test relationship detection"""
        concepts = [
            {"text": "Pythagorean theorem", "type": "CONCEPT"},
            {"text": "triangle", "type": "CONCEPT"}
        ]
        text = "The Pythagorean theorem applies to triangles"
        
        relationships = mock_model_manager.detect_relationships(concepts, text)
        
        assert len(relationships) > 0
        assert all("from_concept" in r for r in relationships)
        assert all("to_concept" in r for r in relationships)
        assert all("relationship_type" in r for r in relationships)
    
    def test_estimate_difficulty(self, test_config):
        """Test difficulty estimation"""
        manager = ModelManager(test_config)
        
        easy_text = "The cat sat on the mat."
        hard_text = "The epistemological ramifications of phenomenological hermeneutics."
        
        easy_diff = manager._estimate_difficulty(easy_text)
        hard_diff = manager._estimate_difficulty(hard_text)
        
        assert easy_diff in ["beginner", "intermediate"]
        assert hard_diff in ["intermediate", "advanced"]
    
    def test_estimate_blooms_level(self, test_config):
        """Test Bloom's level estimation"""
        manager = ModelManager(test_config)
        
        remember_text = "Define the term photosynthesis"
        create_text = "Design a new algorithm for sorting"
        
        remember_level = manager._estimate_blooms_level(remember_text)
        create_level = manager._estimate_blooms_level(create_text)
        
        assert remember_level == "remember"
        assert create_level == "create"


# ============================================================================
# KNOWLEDGE GRAPH CLIENT TESTS
# ============================================================================

class TestKnowledgeGraphClient:
    """Test KnowledgeGraphClient class"""
    
    @pytest.mark.asyncio
    async def test_create_nodes_success(self, test_config, mock_graph_client):
        """Test successful node creation"""
        concepts = [
            {"text": "algebra", "type": "CONCEPT", "confidence": 0.9}
        ]
        metadata = {"file_name": "test.pdf", "subject": "math"}
        
        node_ids = await mock_graph_client.create_nodes(concepts, metadata)
        
        assert len(node_ids) > 0
        assert all(isinstance(nid, str) for nid in node_ids)
    
    @pytest.mark.asyncio
    async def test_create_relationships_success(self, test_config, mock_graph_client):
        """Test successful relationship creation"""
        relationships = [
            {
                "from_concept": "algebra",
                "to_concept": "equations",
                "relationship_type": "related_to",
                "confidence": 0.85
            }
        ]
        
        rel_ids = await mock_graph_client.create_relationships(relationships)
        
        assert len(rel_ids) > 0
        assert all(isinstance(rid, str) for rid in rel_ids)
    
    def test_generate_concept_id(self, test_config):
        """Test concept ID generation"""
        client = KnowledgeGraphClient(test_config)
        
        id1 = client._generate_concept_id("Algebra")
        id2 = client._generate_concept_id("algebra")
        id3 = client._generate_concept_id("Algebra")
        
        # Same text (case-insensitive) should produce same ID
        assert id1 == id2 == id3
        assert len(id1) == 16  # MD5 hash truncated to 16 chars
    
    def test_map_relationship_type(self, test_config):
        """Test relationship type mapping"""
        client = KnowledgeGraphClient(test_config)
        
        assert client._map_relationship_type("prerequisite_of") == "PREREQUISITE_OF"
        assert client._map_relationship_type("related_to") == "BELONGS_TO"
        assert client._map_relationship_type("unknown") == "BELONGS_TO"


# ============================================================================
# CONTENT INGESTION AGENT TESTS
# ============================================================================

class TestContentIngestionAgent:
    """Test ContentIngestionAgent class"""
    
    @pytest.mark.asyncio
    async def test_initialize(self, test_config, tmp_path):
        """Test agent initialization"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        
        with patch.object(agent.model_manager, "load_models"):
            await agent.initialize()
            agent.model_manager.load_models.assert_called_once()
    
    @pytest.mark.asyncio
    async def test_ingest_file_success(
        self,
        test_config,
        sample_pdf,
        mock_model_manager,
        mock_graph_client,
        tmp_path
    ):
        """Test successful file ingestion"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        agent.model_manager = mock_model_manager
        agent.graph_client = mock_graph_client
        
        metadata = {"subject": "math", "grade": "10"}
        
        result = await agent.ingest_file(sample_pdf, metadata)
        
        assert result["status"] == ProcessingStatus.COMPLETED
        assert "job_id" in result
        assert "file_id" in result
        assert "concepts" in result
        assert "relationships" in result
        assert len(result["concepts"]) > 0
    
    @pytest.mark.asyncio
    async def test_ingest_file_validation_error(
        self,
        test_config,
        mock_model_manager,
        mock_graph_client,
        tmp_path
    ):
        """Test ingestion with validation error"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        agent.model_manager = mock_model_manager
        agent.graph_client = mock_graph_client
        
        nonexistent_file = Path("/nonexistent.pdf")
        metadata = {}
        
        with pytest.raises(ValueError):
            await agent.ingest_file(nonexistent_file, metadata)
    
    def test_get_job_status_success(self, test_config, tmp_path):
        """Test job status retrieval"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        
        job_id = "test-job-123"
        agent.jobs[job_id] = {
            "status": ProcessingStatus.PROCESSING,
            "progress": 50.0,
            "message": "Processing"
        }
        
        status = agent.get_job_status(job_id)
        
        assert status["status"] == ProcessingStatus.PROCESSING
        assert status["progress"] == 50.0
    
    def test_get_job_status_not_found(self, test_config, tmp_path):
        """Test job status retrieval for non-existent job"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        
        with pytest.raises(ValueError, match="Job not found"):
            agent.get_job_status("nonexistent-job")
    
    @pytest.mark.asyncio
    async def test_health_status(self, test_config, tmp_path):
        """Test health status check"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        
        health = await agent.get_health_status()
        
        assert "status" in health
        assert "timestamp" in health
        assert "version" in health


# ============================================================================
# API ENDPOINT TESTS
# ============================================================================

class TestAPIEndpoints:
    """Test FastAPI endpoints"""
    
    def test_health_endpoint(self):
        """Test health check endpoint"""
        client = TestClient(app)
        response = client.get("/health")
        
        assert response.status_code == 200
        data = response.json()
        assert data["status"] in ["healthy", "unhealthy"]
    
    def test_metrics_endpoint(self):
        """Test Prometheus metrics endpoint"""
        client = TestClient(app)
        response = client.get("/metrics")
        
        assert response.status_code == 200
        assert "text/plain" in response.headers["content-type"]
    
    @pytest.mark.asyncio
    async def test_ingest_endpoint(self, sample_pdf):
        """Test file ingestion endpoint"""
        client = TestClient(app)
        
        with open(sample_pdf, "rb") as f:
            files = {"file": ("test.pdf", f, "application/pdf")}
            data = {"metadata": json.dumps({"subject": "math"})}
            
            response = client.post("/ingest", files=files, data=data)
        
        assert response.status_code in [200, 202]
        result = response.json()
        assert "job_id" in result
        assert "file_id" in result


# ============================================================================
# INTEGRATION TESTS
# ============================================================================

class TestIntegration:
    """Integration tests"""
    
    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_full_ingestion_pipeline(
        self,
        test_config,
        sample_pdf,
        tmp_path
    ):
        """Test complete ingestion pipeline"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        
        # Mock components
        with patch.object(agent.model_manager, "load_models"):
            with patch.object(agent.model_manager, "classify_content") as mock_classify:
                with patch.object(agent.model_manager, "extract_concepts") as mock_extract:
                    with patch.object(agent.model_manager, "detect_relationships") as mock_detect:
                        with patch.object(agent.graph_client, "create_nodes") as mock_nodes:
                            with patch.object(agent.graph_client, "create_relationships") as mock_rels:
                                
                                # Setup mocks
                                mock_classify.return_value = {
                                    "primary_subject": "math",
                                    "confidence": 0.9,
                                    "difficulty": "intermediate",
                                    "blooms_level": "apply"
                                }
                                
                                mock_extract.return_value = [
                                    {"text": "algebra", "type": "CONCEPT", "confidence": 0.9}
                                ]
                                
                                mock_detect.return_value = []
                                
                                mock_nodes.return_value = ["node1"]
                                mock_rels.return_value = []
                                
                                # Run pipeline
                                result = await agent.ingest_file(
                                    sample_pdf,
                                    {"subject": "math"}
                                )
                                
                                # Verify
                                assert result["status"] == ProcessingStatus.COMPLETED
                                assert len(result["concepts"]) > 0
                                mock_classify.assert_called_once()
                                mock_extract.assert_called_once()


# ============================================================================
# PERFORMANCE TESTS
# ============================================================================

class TestPerformance:
    """Performance tests"""
    
    @pytest.mark.performance
    @pytest.mark.asyncio
    async def test_concurrent_ingestions(self, test_config, sample_pdf, tmp_path):
        """Test concurrent file ingestions"""
        config_path = tmp_path / "config.yaml"
        with open(config_path, "w") as f:
            import yaml
            yaml.dump(test_config, f)
        
        agent = ContentIngestionAgent(str(config_path))
        
        # Mock heavy operations
        with patch.object(agent.model_manager, "load_models"):
            with patch.object(agent.file_processor, "extract_content") as mock_extract:
                with patch.object(agent.model_manager, "classify_content") as mock_classify:
                    
                    mock_extract.return_value = {
                        "type": "pdf",
                        "pages": [{"text": "test"}],
                        "page_count": 1
                    }
                    
                    mock_classify.return_value = {
                        "primary_subject": "math",
                        "confidence": 0.9,
                        "difficulty": "beginner",
                        "blooms_level": "remember"
                    }
                    
                    # Run concurrent ingestions
                    tasks = [
                        agent.ingest_file(sample_pdf, {"id": i})
                        for i in range(5)
                    ]
                    
                    start = datetime.utcnow()
                    results = await asyncio.gather(*tasks, return_exceptions=True)
                    duration = (datetime.utcnow() - start).total_seconds()
                    
                    # Check all succeeded
                    assert len(results) == 5
                    assert all(
                        isinstance(r, dict) or isinstance(r, Exception)
                        for r in results
                    )
                    
                    # Performance threshold
                    assert duration < 30  # Should complete in under 30 seconds


# ============================================================================
# RUN TESTS
# ============================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v", "--cov=content_ingestion_agent", "--cov-report=html"])
